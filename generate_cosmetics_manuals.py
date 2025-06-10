#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
化妆品用户手册和销量数据生成器
生成三个化妆品品牌的用户手册和销量分析报告
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import random
import numpy as np
from datetime import datetime, timedelta
import calendar
import pandas as pd

def create_styled_document(title):
    """创建带有基本样式的文档"""
    doc = Document()
    
    # 设置文档标题
    title_paragraph = doc.add_heading(title, 0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置字体
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    return doc

def generate_yingfeifan_manual():
    """生成赢飞凡化妆品用户手册"""
    print("📝 生成赢飞凡用户手册...")
    
    doc = create_styled_document("赢飞凡™ 奢华抗衰精华液用户手册")
    
    # 产品概述
    doc.add_heading("产品概述", level=1)
    doc.add_paragraph("赢飞凡™ 奢华抗衰精华液是一款革命性的高端护肤产品，专为追求卓越护肤效果的现代女性设计。")
    doc.add_paragraph("🌟 品牌理念：赢在起点，飞越时光，凡而不凡")
    doc.add_paragraph("💎 产品定位：奢华抗衰，科技护肤的典范")
    
    # 核心成分
    doc.add_heading("核心成分与功效", level=1)
    doc.add_paragraph("本产品采用独家专利配方，融合多种珍贵活性成分：")
    doc.add_paragraph("• 六胜肽复合物（15%）：深层抗皱，重塑肌肤弹性")
    doc.add_paragraph("• 烟酰胺（5%）：提亮肌肤，改善暗沉")
    doc.add_paragraph("• 透明质酸钠（3%）：深度保湿，锁水24小时")
    doc.add_paragraph("• 维生素C衍生物（2%）：抗氧化，促进胶原蛋白合成")
    doc.add_paragraph("• 白藜芦醇（1%）：抗衰老，延缓肌肤老化")
    
    # 适用人群
    doc.add_heading("适用人群", level=1)
    doc.add_paragraph("✅ 25-45岁关注抗衰老的女性")
    doc.add_paragraph("✅ 肌肤出现细纹、松弛迹象的用户")
    doc.add_paragraph("✅ 追求高端护肤体验的消费者")
    doc.add_paragraph("✅ 适合所有肌肤类型，包括敏感肌")
    
    # 使用方法
    doc.add_heading("使用方法", level=1)
    doc.add_paragraph("🌅 晨间护理：")
    doc.add_paragraph("1. 洁面后，取2-3滴精华液于掌心")
    doc.add_paragraph("2. 轻拍至面部和颈部，避开眼周")
    doc.add_paragraph("3. 轻柔按摩至完全吸收")
    doc.add_paragraph("4. 后续使用面霜和防晒")
    
    doc.add_paragraph("🌙 夜间护理：")
    doc.add_paragraph("1. 晚间洁面后使用")
    doc.add_paragraph("2. 可适量增加用量（3-4滴）")
    doc.add_paragraph("3. 配合按摩手法，促进吸收")
    doc.add_paragraph("4. 后续使用夜间面霜")
    
    # 注意事项
    doc.add_heading("注意事项", level=1)
    doc.add_paragraph("⚠️ 首次使用前请进行过敏测试")
    doc.add_paragraph("⚠️ 避免接触眼部，如不慎接触请立即用清水冲洗")
    doc.add_paragraph("⚠️ 孕期和哺乳期女性请咨询医生后使用")
    doc.add_paragraph("⚠️ 请存放在阴凉干燥处，避免阳光直射")
    doc.add_paragraph("⚠️ 开封后请在6个月内使用完毕")
    
    # 产品规格
    doc.add_heading("产品规格", level=1)
    doc.add_paragraph("📦 容量：30ml")
    doc.add_paragraph("💰 建议零售价：¥1,299")
    doc.add_paragraph("🏭 生产商：赢飞凡生物科技有限公司")
    doc.add_paragraph("📍 产地：中国上海")
    doc.add_paragraph("📅 保质期：3年（未开封）")
    
    doc.save('documents/real/赢飞凡用户手册.docx')
    print("✅ 赢飞凡用户手册生成完成")

def generate_kangshuiqi_manual():
    """生成康水期化妆品用户手册"""
    print("📝 生成康水期用户手册...")
    
    doc = create_styled_document("康水期™ 舒缓修护面膜用户手册")
    
    # 产品概述
    doc.add_heading("产品概述", level=1)
    doc.add_paragraph("康水期™ 舒缓修护面膜是专为敏感肌和问题肌肤研发的温和护理产品。")
    doc.add_paragraph("🌿 品牌理念：康复肌肤，水润如期，期待新生")
    doc.add_paragraph("🍃 产品定位：天然温和，专业修护")
    
    # 核心成分
    doc.add_heading("核心成分与功效", level=1)
    doc.add_paragraph("精选天然植物提取物，温和有效：")
    doc.add_paragraph("• 积雪草提取物（10%）：舒缓镇静，修护受损肌肤")
    doc.add_paragraph("• 神经酰胺（8%）：重建肌肤屏障，锁住水分")
    doc.add_paragraph("• 透明质酸（5%）：深层补水，持久保湿")
    doc.add_paragraph("• 甘草酸二钾（3%）：抗炎舒缓，减少红肿")
    doc.add_paragraph("• 尿囊素（2%）：促进细胞再生，加速修护")
    doc.add_paragraph("• 马齿苋提取物（1%）：天然抗菌，净化肌肤")
    
    # 适用人群
    doc.add_heading("适用人群", level=1)
    doc.add_paragraph("✅ 敏感肌肤用户")
    doc.add_paragraph("✅ 肌肤屏障受损人群")
    doc.add_paragraph("✅ 经常化妆需要深度清洁修护的用户")
    doc.add_paragraph("✅ 换季时肌肤不稳定的人群")
    doc.add_paragraph("✅ 医美术后需要修护的用户")
    
    # 使用方法
    doc.add_heading("使用方法", level=1)
    doc.add_paragraph("💧 基础护理（每周2-3次）：")
    doc.add_paragraph("1. 彻底清洁面部，用毛巾轻拍至半干")
    doc.add_paragraph("2. 撕开面膜包装，展开面膜")
    doc.add_paragraph("3. 贴合面部轮廓，避开眼唇部位")
    doc.add_paragraph("4. 静敷15-20分钟")
    doc.add_paragraph("5. 撕下面膜，轻拍剩余精华至吸收")
    
    doc.add_paragraph("🆘 急救护理（肌肤敏感时）：")
    doc.add_paragraph("1. 可连续使用3-5天")
    doc.add_paragraph("2. 每次使用时间可延长至25分钟")
    doc.add_paragraph("3. 使用后无需清洗，直接进行后续护理")
    
    # 注意事项
    doc.add_heading("注意事项", level=1)
    doc.add_paragraph("⚠️ 使用前请确保面部清洁")
    doc.add_paragraph("⚠️ 如出现过敏反应请立即停用")
    doc.add_paragraph("⚠️ 一次性使用，请勿重复使用")
    doc.add_paragraph("⚠️ 请存放在阴凉干燥处")
    doc.add_paragraph("⚠️ 儿童请勿接触")
    
    # 产品规格
    doc.add_heading("产品规格", level=1)
    doc.add_paragraph("📦 规格：25ml×5片装")
    doc.add_paragraph("💰 建议零售价：¥168")
    doc.add_paragraph("🏭 生产商：康水期生物技术有限公司")
    doc.add_paragraph("📍 产地：中国广州")
    doc.add_paragraph("📅 保质期：2年")
    
    doc.save('documents/real/康水期用户手册.docx')
    print("✅ 康水期用户手册生成完成")

def generate_anxintang_manual():
    """生成安心唐化妆品用户手册"""
    print("📝 生成安心唐用户手册...")
    
    doc = create_styled_document("安心唐™ 孕妇专用护肤套装用户手册")
    
    # 产品概述
    doc.add_heading("产品概述", level=1)
    doc.add_paragraph("安心唐™ 孕妇专用护肤套装是专为孕期和哺乳期女性设计的安全护肤产品。")
    doc.add_paragraph("👶 品牌理念：安全护肤，心无旁骛，唐风雅韵")
    doc.add_paragraph("🤱 产品定位：孕期专用，安全第一")
    
    # 套装内容
    doc.add_heading("套装内容", level=1)
    doc.add_paragraph("本套装包含完整的孕期护肤方案：")
    doc.add_paragraph("• 温和洁面乳（120ml）：氨基酸配方，温和清洁")
    doc.add_paragraph("• 保湿爽肤水（150ml）：无酒精配方，深层补水")
    doc.add_paragraph("• 滋润面霜（50ml）：天然保湿，长效滋润")
    doc.add_paragraph("• 妊娠纹预防霜（100ml）：专业配方，预防妊娠纹")
    doc.add_paragraph("• 唇部护理膏（4g）：天然成分，安全可食用级别")
    
    # 核心成分
    doc.add_heading("核心成分与功效", level=1)
    doc.add_paragraph("严格筛选孕期安全成分：")
    doc.add_paragraph("• 燕麦提取物：温和清洁，舒缓敏感")
    doc.add_paragraph("• 玻尿酸钠：安全保湿，不含激素")
    doc.add_paragraph("• 乳木果油：天然滋润，修护肌肤")
    doc.add_paragraph("• 维生素E：抗氧化，预防色斑")
    doc.add_paragraph("• 甘油：锁水保湿，温和无刺激")
    doc.add_paragraph("• 可可脂：预防妊娠纹，增强肌肤弹性")
    
    # 适用人群
    doc.add_heading("适用人群", level=1)
    doc.add_paragraph("✅ 孕期女性（孕早期、中期、晚期均可使用）")
    doc.add_paragraph("✅ 哺乳期女性")
    doc.add_paragraph("✅ 备孕期女性")
    doc.add_paragraph("✅ 敏感肌肤用户")
    doc.add_paragraph("✅ 追求天然护肤的用户")
    
    # 使用方法
    doc.add_heading("使用方法", level=1)
    doc.add_paragraph("🌅 晨间护理流程：")
    doc.add_paragraph("1. 温和洁面乳：取适量，加水起泡，轻柔按摩后清洗")
    doc.add_paragraph("2. 保湿爽肤水：用化妆棉轻拍全脸")
    doc.add_paragraph("3. 滋润面霜：取黄豆大小，均匀涂抹")
    doc.add_paragraph("4. 唇部护理膏：涂抹于唇部")
    
    doc.add_paragraph("🌙 夜间护理流程：")
    doc.add_paragraph("1. 重复晨间护理步骤1-3")
    doc.add_paragraph("2. 妊娠纹预防霜：涂抹于腹部、大腿、臀部等易长纹部位")
    doc.add_paragraph("3. 轻柔按摩至完全吸收")
    
    # 安全保障
    doc.add_heading("安全保障", level=1)
    doc.add_paragraph("🔬 严格的安全标准：")
    doc.add_paragraph("• 通过孕期安全性测试")
    doc.add_paragraph("• 不含酒精、激素、重金属")
    doc.add_paragraph("• 不含孕期禁用成分（如维A酸、水杨酸等）")
    doc.add_paragraph("• 通过敏感性测试")
    doc.add_paragraph("• 获得妇产科专家推荐")
    
    # 注意事项
    doc.add_heading("注意事项", level=1)
    doc.add_paragraph("⚠️ 如有特殊过敏史，请先咨询医生")
    doc.add_paragraph("⚠️ 使用过程中如有不适请立即停用")
    doc.add_paragraph("⚠️ 请存放在儿童接触不到的地方")
    doc.add_paragraph("⚠️ 避免阳光直射，常温保存")
    doc.add_paragraph("⚠️ 开封后请在12个月内使用完毕")
    
    # 产品规格
    doc.add_heading("产品规格", level=1)
    doc.add_paragraph("📦 套装规格：5件套")
    doc.add_paragraph("💰 建议零售价：¥599")
    doc.add_paragraph("🏭 生产商：安心唐母婴用品有限公司")
    doc.add_paragraph("📍 产地：中国杭州")
    doc.add_paragraph("📅 保质期：3年")
    
    doc.save('documents/real/安心唐用户手册.docx')
    print("✅ 安心唐用户手册生成完成")

def generate_monthly_data(base_value, years, seasonal_pattern, covid_impact, growth_trend, brand_name, all_brands_data=None):
    """生成月度销量数据（考虑品牌间相互影响）"""
    data = []
    current_value = base_value
    
    # 定义每个品牌的定价策略
    pricing_strategies = {
        '赢飞凡': {
            'base_price': 100,
            'price_changes': {
                2018: 1.0, 2019: 1.05, 2020: 0.95, 2021: 1.02, 
                2022: 1.08, 2023: 1.12, 2024: 1.15, 2025: 1.18
            }
        },
        '康水期': {
            'base_price': 85,
            'price_changes': {
                2018: 1.0, 2019: 1.03, 2020: 1.08, 2021: 1.12, 
                2022: 1.06, 2023: 1.10, 2024: 1.14, 2025: 1.17
            }
        },
        '安心唐': {
            'base_price': 120,
            'price_changes': {
                2018: 1.0, 2019: 1.02, 2020: 1.15, 2021: 1.18, 
                2022: 1.10, 2023: 1.13, 2024: 1.16, 2025: 1.20
            }
        }
    }
    
    # 品牌间相互影响系数
    brand_interactions = {
        '赢飞凡': {
            'synergy_with': '康水期',  # 搭配效果好
            'synergy_factor': 1.45,   # 搭配时销量暴涨45%
            'compete_with': None
        },
        '康水期': {
            'synergy_with': '赢飞凡',  # 与赢飞凡搭配效果好
            'synergy_factor': 1.65,   # 搭配时销量狂飙65%
            'compete_with': '安心唐',  # 与安心唐竞争
            'compete_factor': 0.75    # 竞争时销量暴跌25%
        },
        '安心唐': {
            'synergy_with': None,
            'compete_with': '康水期',  # 与康水期竞争
            'compete_factor': 0.68    # 竞争时销量重挫32%
        }
    }
    
    for year in range(2018, 2026):
        # 获取当年的价格系数
        price_factor = pricing_strategies[brand_name]['price_changes'][year]
        base_price = pricing_strategies[brand_name]['base_price']
        current_price = base_price * price_factor
        
        for month in range(1, 13):
            # 基础季节性影响
            seasonal_factor = seasonal_pattern[month - 1]
            
            # 新冠影响 (2020年3月-2022年12月)
            covid_factor = 1.0
            if year == 2020 and month >= 3:
                covid_factor = covid_impact['2020']
            elif year == 2021:
                covid_factor = covid_impact['2021']
            elif year == 2022:
                covid_factor = covid_impact['2022']
            
            # 年度增长趋势
            if year > 2018:
                current_value *= (1 + growth_trend + random.uniform(-0.02, 0.02))
            
            # 品牌间相互影响（从2020年开始显现）
            interaction_factor = 1.0
            if year >= 2020 and all_brands_data:
                interactions = brand_interactions[brand_name]
                
                # 协同效应（搭配使用）
                if interactions.get('synergy_with'):
                    synergy_brand = interactions['synergy_with']
                    # 模拟搭配购买概率（基于季节和市场成熟度）
                    synergy_probability = min(0.3 + (year - 2020) * 0.05, 0.5)
                    if random.random() < synergy_probability:
                        interaction_factor *= interactions['synergy_factor']
                
                # 竞争效应
                if interactions.get('compete_with'):
                    compete_brand = interactions['compete_with']
                    # 竞争强度随时间增加
                    compete_intensity = min(0.2 + (year - 2020) * 0.03, 0.4)
                    if random.random() < compete_intensity:
                        interaction_factor *= interactions['compete_factor']
            
            # 随机波动
            random_factor = random.uniform(0.9, 1.1)
            
            # 计算最终值
            final_value = int(current_value * seasonal_factor * covid_factor * interaction_factor * random_factor)
            
            # 月度价格微调（季节性促销等）
            monthly_price_adjustment = random.uniform(0.95, 1.05)
            final_price = current_price * monthly_price_adjustment
            
            data.append({
                'year': year,
                'month': month,
                'sales_volume': final_value,
                'sales_amount': final_value * final_price,
                'unit_price': final_price
            })
    
    # 将数据保存为CSV文件
    df = pd.DataFrame(data)
    
    # 确保CSV目录存在
    os.makedirs('csv_data', exist_ok=True)
    
    # 保存CSV文件
    csv_filename = f'csv_data/{brand_name}_销量数据_2018-2025.csv'
    df.to_csv(csv_filename, index=False, encoding='utf-8-sig')
    print(f"📊 {brand_name} 销量数据已保存到: {csv_filename}")
    
    return data

def create_sales_report():
    """创建销量分析报告"""
    print("📊 生成销量分析报告...")
    
    doc = create_styled_document("化妆品品牌销量分析报告（2018-2025）")
    
    # 报告概述
    doc.add_heading("报告概述", level=1)
    doc.add_paragraph("本报告分析了赢飞凡、康水期、安心唐三个化妆品品牌在2018年至2025年期间的销量表现，")
    doc.add_paragraph("包括月度销量数据、销售额统计、市场趋势分析以及未来发展规划。")
    
    # 数据生成参数
    brands_data = {
        '赢飞凡': {
            'base_volume': 8000,
            'seasonal': [0.8, 0.7, 0.9, 1.0, 1.1, 1.0, 0.9, 0.8, 1.0, 1.2, 1.4, 1.3],  # 年末销量高
            'covid_impact': {'2020': 0.6, '2021': 0.8, '2022': 0.9},  # 高端产品受影响较大
            'growth_trend': 0.08  # 年增长8%
        },
        '康水期': {
            'base_volume': 12000,
            'seasonal': [0.9, 0.8, 1.1, 1.2, 1.0, 0.9, 0.8, 0.9, 1.0, 1.1, 1.2, 1.0],  # 春夏销量高
            'covid_impact': {'2020': 0.9, '2021': 1.1, '2022': 1.0},  # 面膜需求增加
            'growth_trend': 0.12  # 年增长12%
        },
        '安心唐': {
            'base_volume': 5000,
            'seasonal': [1.0, 1.0, 1.1, 1.0, 1.2, 1.1, 1.0, 0.9, 1.0, 1.1, 1.0, 0.9],  # 相对稳定
            'covid_impact': {'2020': 1.2, '2021': 1.1, '2022': 1.0},  # 孕妇护肤需求增加
            'growth_trend': 0.15  # 年增长15%
        }
    }
    
    # 为每个品牌生成数据和分析
    for brand_name, params in brands_data.items():
        doc.add_heading(f"{brand_name} 销量分析", level=1)
        
        # 生成数据（考虑品牌间相互影响）
        sales_data = generate_monthly_data(
            params['base_volume'],
            range(2018, 2026),
            params['seasonal'],
            params['covid_impact'],
            params['growth_trend'],
            brand_name,
            brands_data
        )
        
        # 创建数据表格
        doc.add_heading("月度销量数据表", level=2)
        
        # 创建表格
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '年月'
        header_cells[1].text = '销量（件）'
        header_cells[2].text = '销售额（万元）'
        header_cells[3].text = '同比增长率'
        
        # 添加数据（每年显示一行汇总）
        yearly_summary = {}
        for data_point in sales_data:
            year = data_point['year']
            if year not in yearly_summary:
                yearly_summary[year] = {'volume': 0, 'amount': 0}
            yearly_summary[year]['volume'] += data_point['sales_volume']
            yearly_summary[year]['amount'] += data_point['sales_amount']
        
        prev_volume = None
        for year in sorted(yearly_summary.keys()):
            row_cells = table.add_row().cells
            volume = yearly_summary[year]['volume']
            amount = yearly_summary[year]['amount'] / 10000  # 转换为万元
            
            growth_rate = "--"
            if prev_volume:
                growth_rate = f"{((volume - prev_volume) / prev_volume * 100):.1f}%"
            
            row_cells[0].text = f"{year}年"
            row_cells[1].text = f"{volume:,}"
            row_cells[2].text = f"{amount:.1f}"
            row_cells[3].text = growth_rate
            
            prev_volume = volume
        
        # 趋势分析
        doc.add_heading("趋势分析", level=2)
        
        if brand_name == '赢飞凡':
            doc.add_paragraph("📈 整体趋势：")
            doc.add_paragraph("• 2018-2019年：稳定增长期，年增长率约8%")
            doc.add_paragraph("• 2020年：受新冠疫情影响，高端消费下降，销量下滑40%")
            doc.add_paragraph("• 2021-2022年：逐步恢复，线上渠道发力")
            doc.add_paragraph("• 2023-2025年：强劲反弹，年增长率超过10%")
            
            doc.add_paragraph("🔍 季节性特征：")
            doc.add_paragraph("• 年末销量高峰：11-12月销量最佳，受双十一、年终奖影响")
            doc.add_paragraph("• 春节前后低谷：1-2月销量相对较低")
            doc.add_paragraph("• 母亲节、情人节等节日带动销量")
            
            doc.add_paragraph("💰 定价策略分析：")
            doc.add_paragraph("• 2019年：品牌升级，适度提价5%以提升品牌形象")
            doc.add_paragraph("• 2020年：疫情期间降价促销，刺激消费需求")
            doc.add_paragraph("• 2021-2025年：随着品牌价值提升和成本上涨，逐步调整产品定位")
            doc.add_paragraph("• 月度价格波动：根据促销活动和库存情况进行微调")
            
            doc.add_paragraph("🤝 品牌协同效应分析：")
            doc.add_paragraph("• 与康水期形成神级搭配：高端抗衰精华+补水面膜的黄金护肤组合")
            doc.add_paragraph("• 2020年起搭配销售效应爆发，带动销量暴涨45%，效果惊人")
            doc.add_paragraph("• 消费者疯狂追捧：'先补水再抗衰'护肤理念引发购买狂潮")
            doc.add_paragraph("• 套装营销奇迹：节日期间组合装销量翻倍，客单价飙升")
            doc.add_paragraph("• 交叉推荐神效：康水期用户转化率高达80%，用户粘性极强")
            
        elif brand_name == '康水期':
            doc.add_paragraph("📈 整体趋势：")
            doc.add_paragraph("• 2018-2019年：快速增长期，年增长率约12%")
            doc.add_paragraph("• 2020年：疫情期间居家护肤需求增加，销量逆势上涨")
            doc.add_paragraph("• 2021-2025年：持续高增长，成为增长最快的品牌")
            
            doc.add_paragraph("🔍 季节性特征：")
            doc.add_paragraph("• 春夏销量高峰：3-6月面膜需求旺盛")
            doc.add_paragraph("• 换季护肤：9-10月销量回升")
            doc.add_paragraph("• 冬季相对平稳：保湿需求稳定")
            
            doc.add_paragraph("💰 定价策略分析：")
            doc.add_paragraph("• 2019年：产品优化，小幅提价3%以覆盖成本")
            doc.add_paragraph("• 2020-2021年：疫情期间需求增长，适度提价8-12%")
            doc.add_paragraph("• 2022年：市场竞争加剧，价格策略趋于保守")
            doc.add_paragraph("• 2023-2025年：品牌力增强，稳步提升产品价值")
            doc.add_paragraph("• 季节性调价：夏季面膜旺季适度提价，冬季促销回馈")
            
            doc.add_paragraph("🤝 品牌协同与竞争分析：")
            doc.add_paragraph("• 与赢飞凡协同效应：作为黄金补水基础，搭配销售狂飙65%，效果震撼")
            doc.add_paragraph("• 协同营销奇迹：'补水+抗衰'套装成为现象级爆款组合")
            doc.add_paragraph("• 与安心唐激烈竞争：同类基础护肤品类，用户争夺白热化")
            doc.add_paragraph("• 竞争重创：2020-2022年受安心唐专业化冲击，销量暴跌25%")
            doc.add_paragraph("• 差异化血战：通过大众化路线与安心唐专业化展开殊死较量")
            doc.add_paragraph("• 渠道突围：全力布局线上平台，避开与安心唐的正面血拼")
            
        else:  # 安心唐
            doc.add_paragraph("📈 整体趋势：")
            doc.add_paragraph("• 2018-2019年：市场培育期，增长率约15%")
            doc.add_paragraph("• 2020年：疫情期间孕妇护肤意识提升，销量大增20%")
            doc.add_paragraph("• 2021-2025年：三胎政策推动，持续高增长")
            
            doc.add_paragraph("🔍 季节性特征：")
            doc.add_paragraph("• 相对稳定：孕妇护肤需求不受季节影响明显")
            doc.add_paragraph("• 5月母亲节、10月备孕高峰期销量略高")
            doc.add_paragraph("• 年末相对平稳：节日消费对孕妇产品影响较小")
            
            doc.add_paragraph("💰 定价策略分析：")
            doc.add_paragraph("• 2019年：市场教育期，保持价格稳定以建立用户信任")
            doc.add_paragraph("• 2020-2021年：疫情推动孕妇护肤意识，提价15-18%")
            doc.add_paragraph("• 2022年：三胎政策利好，适度调整价格策略")
            doc.add_paragraph("• 2023-2025年：专业化定位，持续提升产品价值")
            doc.add_paragraph("• 节日营销：母亲节等特殊时期推出限量套装")
            
            doc.add_paragraph("⚔️ 市场竞争分析：")
            doc.add_paragraph("• 与康水期血腥竞争：基础护肤品类用户争夺进入白热化阶段")
            doc.add_paragraph("• 竞争重创：受康水期大众化策略疯狂冲击，销量重挫32%")
            doc.add_paragraph("• 差异化救命稻草：专业孕妇护肤定位成为唯一突围武器")
            doc.add_paragraph("• 绝地反击：疯狂强化专业认证，与妇产科医院深度绑定")
            doc.add_paragraph("• 细分市场死守：拼死专注孕期护肤，誓死避免价格血战")
            doc.add_paragraph("• 品牌生死线：专业性和安全性成为最后的用户忠诚度防线")
    
    # 综合分析
    doc.add_heading("综合市场分析", level=1)
    
    doc.add_heading("市场表现对比", level=2)
    doc.add_paragraph("📊 增长率排名：")
    doc.add_paragraph("1. 安心唐：年均增长15%，细分市场领导者")
    doc.add_paragraph("2. 康水期：年均增长12%，大众市场明星产品")
    doc.add_paragraph("3. 赢飞凡：年均增长8%，高端市场稳定增长")
    
    doc.add_heading("新冠疫情影响分析", level=2)
    doc.add_paragraph("🦠 疫情对不同品牌的差异化影响：")
    doc.add_paragraph("• 赢飞凡：高端消费受冲击，2020年销量下降40%")
    doc.add_paragraph("• 康水期：居家护肤需求增加，面膜销量逆势增长10%")
    doc.add_paragraph("• 安心唐：孕妇护肤意识提升，销量增长20%")
    
    doc.add_heading("品牌间相互影响分析", level=2)
    doc.add_paragraph("🔗 品牌神级协同效应：")
    doc.add_paragraph("• 赢飞凡+康水期：形成'补水+抗衰'无敌护肤方案，效果爆表")
    doc.add_paragraph("• 协同销售奇迹：2020年起消费者护肤意识爆发，搭配购买狂潮来袭")
    doc.add_paragraph("• 互补性神效：不同价位和功效完美契合，消费升级需求暴涨")
    doc.add_paragraph("• 交叉营销奇迹：通过一个品牌疯狂带动另一品牌销量飙升45-65%")

    doc.add_paragraph("⚔️ 品牌血腥竞争关系：")
    doc.add_paragraph("• 康水期vs安心唐：基础护肤品类的殊死较量，战况惨烈")
    doc.add_paragraph("• 竞争白热化时期：2020年后市场厮杀激烈，用户争夺进入血战模式")
    doc.add_paragraph("• 差异化生死战：康水期大众化突围，安心唐专业化死守阵地")
    doc.add_paragraph("• 市场分化求生：竞争推动各品牌向细分领域疯狂深耕，避免正面血拼")
    
    doc.add_heading("定价策略演变分析", level=2)
    doc.add_paragraph("💰 各品牌定价策略的周期性特征：")
    doc.add_paragraph("• 市场培育期：新品牌通过稳定价格建立市场信任")
    doc.add_paragraph("• 增长期：随着品牌认知度提升，适度调整价格体现价值")
    doc.add_paragraph("• 成熟期：基于成本变化和竞争环境灵活调价")
    doc.add_paragraph("• 特殊时期：疫情等外部因素影响下的价格策略调整")
    doc.add_paragraph("• 季节性调价：根据需求旺季和促销节点进行价格优化")
    doc.add_paragraph("\n📊 价格变动的主要驱动因素：")
    doc.add_paragraph("• 原材料成本波动：影响基础定价策略")
    doc.add_paragraph("• 品牌价值提升：支撑产品溢价能力")
    doc.add_paragraph("• 市场竞争态势：决定价格调整空间")
    doc.add_paragraph("• 消费者接受度：影响价格策略的执行效果")
    doc.add_paragraph("• 渠道政策变化：促销活动和渠道返利影响终端价格")
    doc.add_paragraph("• 品牌间相互影响：协同效应和竞争关系对定价的影响")
    
    doc.add_heading("市场机遇与挑战", level=2)
    doc.add_paragraph("🎯 机遇：")
    doc.add_paragraph("• 消费升级趋势持续")
    doc.add_paragraph("• 线上渠道快速发展")
    doc.add_paragraph("• 细分市场需求增长")
    doc.add_paragraph("• 国货品牌认知度提升")
    
    doc.add_paragraph("⚠️ 挑战：")
    doc.add_paragraph("• 市场竞争加剧")
    doc.add_paragraph("• 原材料成本上涨")
    doc.add_paragraph("• 监管政策趋严")
    doc.add_paragraph("• 消费者需求多样化")
    
    # 未来规划
    doc.add_heading("未来发展规划（2026-2028）", level=1)
    
    doc.add_heading("赢飞凡发展规划", level=2)
    doc.add_paragraph("🎯 战略目标：巩固高端市场地位，年销量突破20万件")
    doc.add_paragraph("📋 具体措施：")
    doc.add_paragraph("• 产品升级：推出超浓缩精华液系列")
    doc.add_paragraph("• 渠道拓展：进入高端百货和免税店")
    doc.add_paragraph("• 品牌建设：邀请国际明星代言")
    doc.add_paragraph("• 技术创新：投资研发新型抗衰成分")
    doc.add_paragraph("• 协同营销：深化与康水期的搭配销售策略")
    doc.add_paragraph("• 套装开发：推出'补水+抗衰'组合装，提升客单价")
    
    doc.add_heading("康水期发展规划", level=2)
    doc.add_paragraph("🎯 战略目标：成为面膜品类领导品牌，年销量突破50万件")
    doc.add_paragraph("📋 具体措施：")
    doc.add_paragraph("• 产品线扩展：推出不同功效的面膜系列")
    doc.add_paragraph("• 渠道下沉：进入三四线城市市场")
    doc.add_paragraph("• 营销创新：KOL合作和社交媒体营销")
    doc.add_paragraph("• 供应链优化：建设自有生产基地")
    doc.add_paragraph("• 协同发展：与赢飞凡联合推广'护肤三步法'理念")
    doc.add_paragraph("• 差异化竞争：避开与安心唐的直接冲突，专注大众市场")
    doc.add_paragraph("• 渠道策略：重点布局线上平台，与安心唐形成渠道差异")
    
    doc.add_heading("安心唐发展规划", level=2)
    doc.add_paragraph("🎯 战略目标：扩展母婴护肤全品类，年销量突破15万件")
    doc.add_paragraph("📋 具体措施：")
    doc.add_paragraph("• 产品扩展：推出婴幼儿护肤系列")
    doc.add_paragraph("• 专业认证：获得更多医院和专家推荐")
    doc.add_paragraph("• 教育营销：开展孕期护肤知识普及")
    doc.add_paragraph("• 渠道合作：与母婴店和医院合作")
    doc.add_paragraph("• 竞争应对：强化专业化定位，避免与康水期价格战")
    doc.add_paragraph("• 护城河建设：通过医学认证和专家背书提升品牌壁垒")
    doc.add_paragraph("• 细分深耕：专注孕产妇群体，建立用户生命周期管理")
    
    doc.add_heading("整体预期", level=2)
    doc.add_paragraph("📈 2026-2028年预期：")
    doc.add_paragraph("• 三个品牌总销量预计突破85万件")
    doc.add_paragraph("• 总销售额预计达到8亿元")
    doc.add_paragraph("• 市场份额在各自细分领域均进入前三")
    doc.add_paragraph("• 品牌价值和消费者认知度显著提升")
    
    doc.save('documents/real/化妆品销量分析报告.docx')
    print("✅ 销量分析报告生成完成")

def generate_all_manuals():
    """生成所有化妆品手册和报告"""
    print("🚀 开始生成化妆品用户手册和销量报告...")
    
    # 确保目录存在
    os.makedirs('documents/real', exist_ok=True)
    
    # 生成用户手册
    generate_yingfeifan_manual()
    generate_kangshuiqi_manual()
    generate_anxintang_manual()
    
    # 生成销量分析报告
    create_sales_report()
    
    print("\n📊 文档生成统计：")
    print("用户手册：3个")
    print("销量报告：1个")
    print("CSV数据文件：3个")
    print("总计：4个docx文件 + 3个csv文件")
    print("\n📁 文件位置：")
    print("  Word文档：documents/real/")
    print("  CSV数据：csv_data/")
    print("\n🎮 化妆品文档和数据文件已准备就绪！")
    print("\n📈 CSV文件包含详细的月度销量数据，可用于进一步分析：")
    print("  • 赢飞凡_销量数据_2018-2025.csv")
    print("  • 康水期_销量数据_2018-2025.csv")
    print("  • 安心唐_销量数据_2018-2025.csv")
    print("\n💡 CSV文件字段说明：")
    print("  • year: 年份")
    print("  • month: 月份")
    print("  • sales_volume: 销量（件）")
    print("  • sales_amount: 销售额（元）")
    print("  • unit_price: 单价（元）")

if __name__ == "__main__":
    generate_all_manuals()