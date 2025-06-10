"""文档生成脚本

生成RAG竞技场游戏所需的真实文档和混淆文档
所有文档以docx格式保存
"""

import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def create_document_with_style(title: str) -> Document:
    """创建带有基本样式的文档"""
    doc = Document()
    
    # 添加标题
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

def generate_real_documents():
    """生成真实的专业文档"""
    
    # 创建documents目录
    os.makedirs('documents/real', exist_ok=True)
    
    # 文档1 - 化妆品成分安全性分析报告
    doc1 = Document()
    doc1.add_heading("化妆品成分安全性分析报告", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc1.add_heading("前言", level=1)
    doc1.add_paragraph("随着消费者对化妆品安全性要求的不断提高，成分安全性评估已成为产品开发的核心环节。本报告基于最新的科学研究和监管要求，对常见化妆品成分进行全面的安全性分析，为消费者选择和企业研发提供科学依据。")
    
    doc1.add_heading("第一章 玻尿酸（透明质酸）安全性分析", level=1)
    doc1.add_paragraph("玻尿酸（Hyaluronic Acid，简称HA）是一种天然存在于人体的糖胺聚糖，由D-葡萄糖醛酸和N-乙酰葡糖胺通过β-1,4和β-1,3糖苷键交替连接而成。其分子结构的独特性赋予了它卓越的保湿性能和生物相容性。")
    
    doc1.add_heading("1.1 分子特性与生物学功能", level=2)
    doc1.add_paragraph("• 分子量范围：1,000-2,000,000 Da（不同分子量具有不同的渗透性和功效）")
    doc1.add_paragraph("• 化学式：(C14H21NO11)n")
    doc1.add_paragraph("• 安全性等级：GRAS（Generally Recognized as Safe）")
    doc1.add_paragraph("• 生物降解性：可被人体透明质酸酶完全降解")
    doc1.add_paragraph("• 致敏性：极低，适合敏感肌肤使用")
    doc1.add_paragraph("• 推荐浓度：0.1%-2.0%（根据产品类型和目标效果调整）")
    
    doc1.add_heading("1.2 透明质酸钠的优势特性", level=2)
    doc1.add_paragraph("透明质酸钠（Sodium Hyaluronate）是玻尿酸的钠盐形式，在化妆品应用中表现出更优异的性能：")
    doc1.add_paragraph("• 分子量更小（通常<300kDa），渗透性显著增强")
    doc1.add_paragraph("• 保湿能力：理论上可结合自身重量1000倍的水分")
    doc1.add_paragraph("• pH稳定范围：6.0-8.0（适合大多数化妆品配方）")
    doc1.add_paragraph("• 与其他成分配伍性良好，无明显配伍禁忌")
    doc1.add_paragraph("• 热稳定性：在80°C以下保持稳定")
    
    doc1.add_heading("1.3 敏感肌适用性评估", level=2)
    doc1.add_paragraph("基于多项临床试验数据，玻尿酸在敏感肌群体中的安全性表现：")
    doc1.add_paragraph("• 刺激性测试：0/100受试者出现不良反应（24小时贴片测试）")
    doc1.add_paragraph("• 致敏性测试：阴性率99.8%（HRIPT测试，n=200）")
    doc1.add_paragraph("• 耐受性：优秀，可长期使用无累积毒性")
    doc1.add_paragraph("• 光敏性：无光敏反应报告")
    doc1.add_paragraph("• 眼部安全性：通过眼部刺激性测试")
    
    doc1.add_heading("1.4 不同分子量玻尿酸的功效差异", level=2)
    doc1.add_paragraph("• 高分子量HA（>1000kDa）：主要在皮肤表面形成保湿膜，即时保湿效果显著")
    doc1.add_paragraph("• 中分子量HA（100-1000kDa）：可渗透至角质层，提供持久保湿")
    doc1.add_paragraph("• 低分子量HA（<100kDa）：可深入真皮层，促进胶原蛋白合成")
    doc1.add_paragraph("• 寡聚透明质酸（<10kDa）：具有抗炎和促进伤口愈合的作用")
    
    doc1.add_paragraph("\n💡 详细的分子量选择指南请参考文档3第2.3节的配方设计原则。")
    
    doc1.save('documents/real/文档1.docx')
    
    # 文档2 - 敏感肌肤护理专业指南
    doc2 = Document()
    doc2.add_heading("敏感肌肤护理专业指南", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc2.add_heading("概述", level=1)
    doc2.add_paragraph("敏感肌肤是一种常见的皮肤状态，影响全球约50%的人群。它不是一种疾病，而是皮肤对外界刺激产生过度反应的表现。正确的护理方法和成分选择对于改善敏感肌肤状况至关重要。本指南基于最新的皮肤科学研究，为敏感肌肤人群提供科学、实用的护理建议。")
    
    doc2.add_heading("第一章 敏感肌肤的科学认知", level=1)
    
    doc2.add_heading("1.1 敏感肌肤特征识别", level=2)
    doc2.add_paragraph("敏感肌肤的主要特征包括：")
    doc2.add_paragraph("• 皮肤屏障功能受损：角质层完整性破坏，TEWL（经皮水分流失）增加")
    doc2.add_paragraph("• 对外界刺激反应强烈：温度变化、化学物质、物理摩擦等")
    doc2.add_paragraph("• 容易出现红肿、瘙痒、刺痛、紧绷感")
    doc2.add_paragraph("• 角质层薄（通常<10μm），毛细血管扩张")
    doc2.add_paragraph("• 皮脂分泌异常，pH值偏碱性（>5.5）")
    doc2.add_paragraph("• 微生物群落失衡，有益菌减少")
    
    doc2.add_heading("1.2 敏感肌肤的分类", level=2)
    doc2.add_paragraph("• 先天性敏感：遗传因素导致，皮肤屏障天生脆弱")
    doc2.add_paragraph("• 后天性敏感：环境因素、不当护理、疾病等引起")
    doc2.add_paragraph("• 暂时性敏感：特定时期（如换季、生理期）出现的敏感状态")
    doc2.add_paragraph("• 永久性敏感：需要长期特殊护理的敏感状态")
    
    doc2.add_heading("第二章 成分选择的科学原则", level=1)
    
    doc2.add_heading("2.1 推荐成分及其机制", level=2)
    doc2.add_paragraph("以下成分经过大量临床验证，对敏感肌肤安全有效：")
    doc2.add_paragraph("• 神经酰胺（Ceramide）：修复皮肤屏障，推荐浓度0.2-1%")
    doc2.add_paragraph("  - 机制：补充角质层脂质，恢复屏障功能")
    doc2.add_paragraph("  - 类型：Ceramide NP、AP、EOP等")
    doc2.add_paragraph("• 烟酰胺（Niacinamide）：抗炎舒缓，推荐浓度2-5%")
    doc2.add_paragraph("  - 机制：抑制炎症介质释放，增强屏障功能")
    doc2.add_paragraph("  - 注意：浓度过高可能引起刺激")
    doc2.add_paragraph("• 透明质酸：温和保湿，推荐分子量500-1800kDa")
    doc2.add_paragraph("  - 机制：形成保湿膜，减少水分流失")
    doc2.add_paragraph("  - 优势：天然成分，致敏性极低")
    doc2.add_paragraph("• 泛醇（Pro-Vitamin B5）：舒缓修复，推荐浓度1-5%")
    doc2.add_paragraph("• 尿囊素（Allantoin）：抗炎镇静，推荐浓度0.1-2%")
    doc2.add_paragraph("• 马齿苋提取物：天然抗炎，推荐浓度0.5-2%")
    
    doc2.add_heading("2.2 应避免的成分", level=2)
    doc2.add_paragraph("以下成分可能加重敏感肌肤的不适症状：")
    doc2.add_paragraph("• 酒精（乙醇）：会破坏皮肤屏障，增加TEWL")
    doc2.add_paragraph("• 人工香精：常见致敏源，特别是柠檬烯、芳樟醇")
    doc2.add_paragraph("• 强刺激性防腐剂：MIT、CMIT、甲醛释放体")
    doc2.add_paragraph("• 果酸类：水杨酸、甘醇酸等去角质成分")
    doc2.add_paragraph("• 挥发性精油：薰衣草、茶树、薄荷等")
    doc2.add_paragraph("• 强碱性成分：皂基清洁剂（pH>8）")
    
    doc2.add_heading("第三章 产品选择与使用指南", level=1)
    
    doc2.add_heading("3.1 面膜选择要点", level=2)
    doc2.add_paragraph("敏感肌肤选择面膜时应注意：")
    doc2.add_paragraph("• 材质选择：无纺布或生物纤维材质，避免无纺布过厚")
    doc2.add_paragraph("• 成分分析：成分表前5位应为温和成分")
    doc2.add_paragraph("• 避免复杂配方：含有多种植物提取物的产品风险较高")
    doc2.add_paragraph("• 品牌选择：优先选择医学护肤品牌或敏感肌专用品牌")
    doc2.add_paragraph("• 安全测试：使用前进行48小时贴片测试")
    doc2.add_paragraph("• 使用频率：每周1-2次，避免过度使用")
    
    doc2.add_paragraph("\n🔍 关于玻尿酸面膜的详细成分分析，请查看文档1第1.3节的敏感肌适用性评估。")
    doc2.add_paragraph("📋 具体的产品配方设计原则，请参考文档5第2章的配方技术要点。")
    
    doc2.save('documents/real/文档2.docx')
    
    # 文档3 - 孕期化妆品使用安全指南
    doc3 = Document()
    doc3.add_heading("孕期化妆品使用安全指南", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc3.add_heading("前言", level=1)
    doc3.add_paragraph("怀孕期间，女性的皮肤状态会发生显著变化，同时对化妆品成分的安全性要求也更加严格。本指南基于国际权威机构的安全评估数据，为孕期女性提供科学、可靠的化妆品使用建议，确保母婴健康的同时维护肌肤状态。")
    
    doc3.add_heading("第一章 孕期皮肤变化与需求", level=1)
    
    doc3.add_heading("1.1 孕期皮肤生理变化", level=2)
    doc3.add_paragraph("怀孕期间，激素水平的剧烈变化会导致多种皮肤问题：")
    doc3.add_paragraph("• 色素沉着增加：雌激素和孕激素刺激黑色素细胞活跃")
    doc3.add_paragraph("• 皮脂分泌异常：可能出现痤疮或皮肤干燥")
    doc3.add_paragraph("• 皮肤敏感性增加：对外界刺激反应更强烈")
    doc3.add_paragraph("• 血管扩张：面部红血丝增多")
    doc3.add_paragraph("• 妊娠纹形成：胶原蛋白和弹性纤维断裂")
    doc3.add_paragraph("• 水肿现象：特别是眼部和面部")
    
    doc3.add_heading("1.2 孕期护肤需求分析", level=2)
    doc3.add_paragraph("• 温和清洁：避免过度清洁破坏皮肤屏障")
    doc3.add_paragraph("• 充分保湿：维持皮肤水油平衡")
    doc3.add_paragraph("• 有效防晒：预防妊娠斑加重")
    doc3.add_paragraph("• 舒缓抗炎：缓解皮肤敏感症状")
    doc3.add_paragraph("• 预防妊娠纹：增强皮肤弹性")
    
    doc3.add_heading("第二章 成分安全性评估", level=1)
    
    doc3.add_heading("2.1 孕期禁用成分（A级风险）", level=2)
    doc3.add_paragraph("以下成分孕期应严格避免，可能对胎儿造成致畸风险：")
    doc3.add_paragraph("• 维A酸及其衍生物：")
    doc3.add_paragraph("  - Tretinoin（全反式维A酸）：FDA妊娠分类C级")
    doc3.add_paragraph("  - Isotretinoin（异维A酸）：FDA妊娠分类X级")
    doc3.add_paragraph("  - Adapalene（阿达帕林）：可能致畸")
    doc3.add_paragraph("• 对苯二酚（Hydroquinone）：")
    doc3.add_paragraph("  - 系统吸收率高达35-45%")
    doc3.add_paragraph("  - 可能影响胎儿发育")
    doc3.add_paragraph("• 某些精油：")
    doc3.add_paragraph("  - 迷迭香、鼠尾草：可能引起子宫收缩")
    doc3.add_paragraph("  - 薄荷、桉树：高浓度使用有风险")
    doc3.add_paragraph("• 甲醛及其释放体：")
    doc3.add_paragraph("  - DMDM Hydantoin、Imidazolidinyl Urea")
    doc3.add_paragraph("  - 可能致癌和致畸")
    
    doc3.add_heading("2.2 谨慎使用成分（B级风险）", level=2)
    doc3.add_paragraph("以下成分需要控制浓度和使用频率：")
    doc3.add_paragraph("• 水杨酸（Salicylic Acid）：")
    doc3.add_paragraph("  - 低浓度（<2%）局部使用相对安全")
    doc3.add_paragraph("  - 避免大面积或高浓度使用")
    doc3.add_paragraph("• 苯甲酸酯类防腐剂：")
    doc3.add_paragraph("  - 可能的内分泌干扰物")
    doc3.add_paragraph("  - 选择不含paraben的产品")
    doc3.add_paragraph("• 化学防晒剂：")
    doc3.add_paragraph("  - Oxybenzone、Octinoxate等")
    doc3.add_paragraph("  - 优先选择物理防晒")
    
    doc3.add_heading("2.3 安全成分推荐（A级安全）", level=2)
    doc3.add_paragraph("孕期可安全使用的成分：")
    doc3.add_paragraph("• 透明质酸：")
    doc3.add_paragraph("  - 分子量大，不易透皮吸收")
    doc3.add_paragraph("  - 安全保湿，无已知副作用")
    doc3.add_paragraph("• 维生素C（L-抗坏血酸）：")
    doc3.add_paragraph("  - 抗氧化，促进胶原合成")
    doc3.add_paragraph("  - 推荐浓度<20%，避免刺激")
    doc3.add_paragraph("• 烟酰胺（维生素B3）：")
    doc3.add_paragraph("  - 改善肌肤状态，控制油脂")
    doc3.add_paragraph("  - 安全性高，推荐浓度2-5%")
    doc3.add_paragraph("• 神经酰胺：")
    doc3.add_paragraph("  - 修复皮肤屏障")
    doc3.add_paragraph("  - 天然成分，安全性极高")
    doc3.add_paragraph("• 甘油：")
    doc3.add_paragraph("  - 基础保湿成分")
    doc3.add_paragraph("  - 使用历史悠久，安全可靠")
    doc3.add_paragraph("• 物理防晒剂：")
    doc3.add_paragraph("  - 氧化锌、二氧化钛")
    doc3.add_paragraph("  - 不被皮肤吸收，安全性最高")
    
    doc3.add_heading("第三章 孕期护肤方案", level=1)
    
    doc3.add_heading("3.1 基础护肤步骤", level=2)
    doc3.add_paragraph("孕期护肤要点：")
    doc3.add_paragraph("• 简化护肤步骤：清洁-保湿-防晒三步骤")
    doc3.add_paragraph("• 选择温和无刺激产品：pH值接近皮肤（5.5-6.5）")
    doc3.add_paragraph("• 避免频繁更换产品：减少过敏风险")
    doc3.add_paragraph("• 注意防晒：选择SPF30+的物理防晒")
    doc3.add_paragraph("• 如有疑虑：及时咨询皮肤科医生")
    
    doc3.add_heading("3.2 特殊问题处理", level=2)
    doc3.add_paragraph("• 妊娠斑预防：")
    doc3.add_paragraph("  - 严格防晒，避免紫外线刺激")
    doc3.add_paragraph("  - 使用维生素C产品淡化色斑")
    doc3.add_paragraph("• 妊娠纹预防：")
    doc3.add_paragraph("  - 使用含有维生素E的身体乳")
    doc3.add_paragraph("  - 适度按摩，促进血液循环")
    doc3.add_paragraph("• 孕期痤疮：")
    doc3.add_paragraph("  - 温和清洁，避免过度去油")
    doc3.add_paragraph("  - 使用低浓度烟酰胺产品")
    
    doc3.add_paragraph("\n🔍 关于维生素C在孕期的详细使用指南，请查看文档4第3.2节的抗氧化成分应用。")
    doc3.add_paragraph("📋 敏感肌孕妇的特殊护理建议，请参考文档2第3章的产品选择指南。")
    
    doc3.save('documents/real/文档3.docx')
    
    # 文档4 - 抗衰老活性成分科学研究报告
    doc4 = Document()
    doc4.add_heading("抗衰老活性成分科学研究报告", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc4.add_heading("研究背景", level=1)
    doc4.add_paragraph("皮肤衰老是一个复杂的生物学过程，涉及内源性衰老（时间衰老）和外源性衰老（光老化）。随着年龄增长，皮肤中胶原蛋白和弹性蛋白逐渐减少，细胞更新速度放缓，抗氧化能力下降。本报告基于近年来的科学研究，系统分析各类抗衰老活性成分的作用机制和临床效果。")
    
    doc4.add_heading("第一章 皮肤衰老的分子机制", level=1)
    
    doc4.add_heading("1.1 内源性衰老机制", level=2)
    doc4.add_paragraph("• 胶原蛋白合成减少：25岁后每年减少1-1.5%")
    doc4.add_paragraph("• 弹性蛋白降解：弹性纤维断裂和变性")
    doc4.add_paragraph("• 细胞更新周期延长：从28天延长至45-60天")
    doc4.add_paragraph("• 透明质酸含量下降：保湿能力减弱")
    doc4.add_paragraph("• 抗氧化酶活性降低：SOD、CAT、GPx等")
    doc4.add_paragraph("• 线粒体功能衰退：ATP产生减少")
    
    doc4.add_heading("1.2 外源性衰老机制", level=2)
    doc4.add_paragraph("• 紫外线损伤：UVA和UVB引起DNA损伤")
    doc4.add_paragraph("• 自由基攻击：ROS和RNS破坏细胞结构")
    doc4.add_paragraph("• 糖化反应：AGEs（糖化终产物）形成")
    doc4.add_paragraph("• 炎症反应：慢性炎症加速衰老")
    doc4.add_paragraph("• 环境污染：PM2.5、重金属等有害物质")
    
    doc4.add_heading("第二章 维A酸类成分研究", level=1)
    
    doc4.add_heading("2.1 维A酸的作用机制", level=2)
    doc4.add_paragraph("维A酸是目前最有效的抗衰老成分之一，其作用机制包括：")
    doc4.add_paragraph("• 促进细胞更新：加速角质细胞脱落，改善肌理")
    doc4.add_paragraph("• 刺激胶原蛋白合成：激活成纤维细胞")
    doc4.add_paragraph("• 减少细纹和皱纹：增加真皮厚度")
    doc4.add_paragraph("• 改善色素沉着：抑制酪氨酸酶活性")
    doc4.add_paragraph("• 调节皮脂分泌：改善毛孔粗大")
    
    doc4.add_heading("2.2 不同类型维A酸比较", level=2)
    doc4.add_paragraph("• 全反式维A酸（Tretinoin）：")
    doc4.add_paragraph("  - 效果最强，刺激性最大")
    doc4.add_paragraph("  - 推荐浓度：0.025%-0.1%")
    doc4.add_paragraph("  - 需要处方，医生指导使用")
    doc4.add_paragraph("• 视黄醇（Retinol）：")
    doc4.add_paragraph("  - 需要转化为维A酸才能发挥作用")
    doc4.add_paragraph("  - 刺激性较小，适合初学者")
    doc4.add_paragraph("  - 推荐浓度：0.1%-1%")
    doc4.add_paragraph("• 视黄醇棕榈酸酯（Retinyl Palmitate）：")
    doc4.add_paragraph("  - 最温和的维A酸衍生物")
    doc4.add_paragraph("  - 效果相对较弱")
    doc4.add_paragraph("  - 适合敏感肌肤")
    
    doc4.add_heading("第三章 多肽类成分研究", level=1)
    
    doc4.add_heading("3.1 多肽分类及作用机制", level=2)
    doc4.add_paragraph("多肽在抗衰老中发挥重要作用：")
    doc4.add_paragraph("• 信号肽（Signal Peptides）：")
    doc4.add_paragraph("  - 刺激胶原蛋白、弹性蛋白产生")
    doc4.add_paragraph("  - 代表成分：Matrixyl（棕榈酰五肽-4）")
    doc4.add_paragraph("  - 临床研究显示皱纹减少68%")
    doc4.add_paragraph("• 载体肽（Carrier Peptides）：")
    doc4.add_paragraph("  - 输送微量元素如铜、锰等")
    doc4.add_paragraph("  - 代表成分：GHK-Cu（蓝铜胜肽）")
    doc4.add_paragraph("  - 促进伤口愈合和组织修复")
    doc4.add_paragraph("• 神经肽（Neurotransmitter Peptides）：")
    doc4.add_paragraph("  - 减少肌肉收缩，类似肉毒素效果")
    doc4.add_paragraph("  - 代表成分：Argireline（乙酰基六肽-8）")
    doc4.add_paragraph("  - 减少表情纹深度30%")
    
    doc4.add_heading("3.2 多肽稳定性与渗透性", level=2)
    doc4.add_paragraph("• 稳定性挑战：多肽易被酶降解")
    doc4.add_paragraph("• 渗透性问题：分子量大，难以透皮")
    doc4.add_paragraph("• 解决方案：脂质体包裹、纳米技术")
    doc4.add_paragraph("• 协同效应：与其他活性成分联合使用")
    
    doc4.add_heading("第四章 抗氧化剂研究", level=1)
    
    doc4.add_heading("4.1 维生素类抗氧化剂", level=2)
    doc4.add_paragraph("• 维生素C（L-抗坏血酸）：")
    doc4.add_paragraph("  - 促进胶原合成，抑制酪氨酸酶")
    doc4.add_paragraph("  - 稳定性差，需要特殊包装")
    doc4.add_paragraph("  - 衍生物：抗坏血酸磷酸镁、抗坏血酸葡糖苷")
    doc4.add_paragraph("• 维生素E（生育酚）：")
    doc4.add_paragraph("  - 保护细胞膜，与维生素C协同作用")
    doc4.add_paragraph("  - 脂溶性，适合干性皮肤")
    doc4.add_paragraph("• 维生素A（视黄醇）：")
    doc4.add_paragraph("  - 既是抗氧化剂又是细胞调节剂")
    doc4.add_paragraph("  - 需要避光保存")
    
    doc4.add_heading("4.2 植物提取物抗氧化剂", level=2)
    doc4.add_paragraph("• 白藜芦醇：")
    doc4.add_paragraph("  - 激活长寿蛋白Sirtuin")
    doc4.add_paragraph("  - 抗炎抗氧化双重作用")
    doc4.add_paragraph("  - 光敏感，需要夜间使用")
    doc4.add_paragraph("• 绿茶提取物（EGCG）：")
    doc4.add_paragraph("  - 多酚类抗氧化剂")
    doc4.add_paragraph("  - 抗炎、抗糖化作用")
    doc4.add_paragraph("  - 浓度推荐：0.5-2%")
    doc4.add_paragraph("• 葡萄籽提取物：")
    doc4.add_paragraph("  - 原花青素含量高")
    doc4.add_paragraph("  - 保护胶原蛋白不被降解")
    
    doc4.add_paragraph("\n🔍 关于孕期使用维生素C的安全性评估，请查看文档3第2.3节的安全成分推荐。")
    doc4.add_paragraph("📋 敏感肌肤使用抗衰老成分的注意事项，请参考文档2第2章的成分选择原则。")
    
    doc4.save('documents/real/文档4.docx')
    
    # 文档5 - 精华液配方分析报告
    doc5 = Document()
    doc5.add_heading("精华液配方分析报告", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc5.add_heading("引言", level=1)
    doc5.add_paragraph("精华液作为现代护肤品中的核心产品，承载着高浓度活性成分的输送任务。其配方设计需要在功效性、稳定性、安全性和感官体验之间找到平衡。本报告深入分析精华液的配方原理、技术要点和质量控制标准，为配方师和品牌方提供专业指导。")
    
    doc5.add_heading("第一章 精华液配方基础理论", level=1)
    
    doc5.add_heading("1.1 配方基础结构", level=2)
    doc5.add_paragraph("精华液的基本配方组成包括：")
    doc5.add_paragraph("• 水相系统（60-90%）：")
    doc5.add_paragraph("  - 去离子水：作为溶剂和载体")
    doc5.add_paragraph("  - 保湿剂：甘油、丁二醇、透明质酸等")
    doc5.add_paragraph("  - 水溶性活性成分：烟酰胺、维生素C等")
    doc5.add_paragraph("• 油相系统（5-20%）：")
    doc5.add_paragraph("  - 植物油：荷荷巴油、角鲨烷等")
    doc5.add_paragraph("  - 脂肪酸：硬脂酸、棕榈酸等")
    doc5.add_paragraph("  - 脂溶性成分：维生素E、植物甾醇等")
    doc5.add_paragraph("• 功能性添加剂（1-10%）：")
    doc5.add_paragraph("  - 乳化剂：聚甘油脂肪酸酯、卵磷脂等")
    doc5.add_paragraph("  - 防腐剂：苯氧乙醇、山梨酸钾等")
    doc5.add_paragraph("  - 调节剂：柠檬酸、氢氧化钠等")
    
    doc5.add_heading("1.2 配方设计原则", level=2)
    doc5.add_paragraph("• 功效导向：明确产品定位和目标功效")
    doc5.add_paragraph("• 成分协同：避免成分间的拮抗作用")
    doc5.add_paragraph("• 稳定性优先：确保产品在货架期内稳定")
    doc5.add_paragraph("• 安全性保障：严格控制刺激性和致敏性")
    doc5.add_paragraph("• 感官优化：追求良好的使用体验")
    doc5.add_paragraph("• 成本控制：在预算范围内实现最佳配方")
    
    doc5.add_heading("第二章 配方技术要点", level=1)
    
    doc5.add_heading("2.1 活性成分配比原则", level=2)
    doc5.add_paragraph("常见活性成分的科学配比：")
    doc5.add_paragraph("• 透明质酸：")
    doc5.add_paragraph("  - 低分子量（<50kDa）：0.05-0.1%，深层保湿")
    doc5.add_paragraph("  - 中分子量（50-1000kDa）：0.1-0.5%，表面保湿")
    doc5.add_paragraph("  - 高分子量（>1000kDa）：0.5-2%，成膜保湿")
    doc5.add_paragraph("• 烟酰胺：")
    doc5.add_paragraph("  - 推荐浓度：2-5%")
    doc5.add_paragraph("  - 与维生素C间隔使用，避免烟酸生成")
    doc5.add_paragraph("  - pH值控制在5.0-7.0")
    doc5.add_paragraph("• 维生素C：")
    doc5.add_paragraph("  - L-抗坏血酸：5-20%，pH<3.5")
    doc5.add_paragraph("  - 抗坏血酸磷酸镁：3-10%，pH 6.0-7.0")
    doc5.add_paragraph("  - 抗坏血酸葡糖苷：2-5%，稳定性好")
    doc5.add_paragraph("• 神经酰胺：")
    doc5.add_paragraph("  - 推荐浓度：0.2-1%")
    doc5.add_paragraph("  - 需要合适的载体系统")
    doc5.add_paragraph("  - 与胆固醇、脂肪酸配合使用")
    doc5.add_paragraph("• 多肽类：")
    doc5.add_paragraph("  - 信号肽：2-5%")
    doc5.add_paragraph("  - 载体肽：1-3%")
    doc5.add_paragraph("  - 神经肽：3-8%")
    
    doc5.add_heading("2.2 pH值控制策略", level=2)
    doc5.add_paragraph("不同成分的pH要求：")
    doc5.add_paragraph("• 酸性环境（pH 3.0-4.5）：")
    doc5.add_paragraph("  - L-抗坏血酸、果酸类")
    doc5.add_paragraph("  - 需要强效缓冲系统")
    doc5.add_paragraph("• 弱酸性环境（pH 4.5-6.5）：")
    doc5.add_paragraph("  - 大多数活性成分的最适pH")
    doc5.add_paragraph("  - 接近皮肤天然pH值")
    doc5.add_paragraph("• 中性环境（pH 6.5-7.5）：")
    doc5.add_paragraph("  - 敏感性成分、蛋白质类")
    doc5.add_paragraph("  - 稳定性较好")
    
    doc5.add_heading("第三章 配方稳定性控制", level=1)
    
    doc5.add_heading("3.1 物理稳定性", level=2)
    doc5.add_paragraph("影响精华液物理稳定性的关键因素：")
    doc5.add_paragraph("• 乳化体系稳定性：")
    doc5.add_paragraph("  - HLB值匹配：选择合适的乳化剂组合")
    doc5.add_paragraph("  - 粒径控制：纳米乳化技术应用")
    doc5.add_paragraph("  - 粘度调节：增稠剂的合理使用")
    doc5.add_paragraph("• 温度稳定性：")
    doc5.add_paragraph("  - 冷热循环测试：-5°C至45°C")
    doc5.add_paragraph("  - 高温加速试验：40°C，75%RH")
    doc5.add_paragraph("• 机械稳定性：")
    doc5.add_paragraph("  - 振动测试：模拟运输条件")
    doc5.add_paragraph("  - 离心测试：3000rpm，30分钟")
    
    doc5.add_heading("3.2 化学稳定性", level=2)
    doc5.add_paragraph("化学稳定性控制措施：")
    doc5.add_paragraph("• 抗氧化系统：")
    doc5.add_paragraph("  - 主抗氧化剂：BHT、BHA、生育酚")
    doc5.add_paragraph("  - 辅助抗氧化剂：抗坏血酸棕榈酸酯")
    doc5.add_paragraph("  - 螯合剂：EDTA、柠檬酸")
    doc5.add_paragraph("• 光稳定性：")
    doc5.add_paragraph("  - 避光包装：深色玻璃瓶、不透明材料")
    doc5.add_paragraph("  - 光稳定剂：二苯酮类、苯并三唑类")
    doc5.add_paragraph("• 水分活度控制：")
    doc5.add_paragraph("  - 目标Aw值：<0.6")
    doc5.add_paragraph("  - 干燥剂使用：硅胶、分子筛")
    
    doc5.add_heading("3.3 微生物稳定性", level=2)
    doc5.add_paragraph("防腐系统设计：")
    doc5.add_paragraph("• 防腐剂选择：")
    doc5.add_paragraph("  - 广谱性：对细菌、真菌、酵母有效")
    doc5.add_paragraph("  - 配伍性：与其他成分相容")
    doc5.add_paragraph("  - 安全性：低刺激、低致敏")
    doc5.add_paragraph("• 防腐增效：")
    doc5.add_paragraph("  - pH调节：酸性环境增强防腐效果")
    doc5.add_paragraph("  - 螯合剂：破坏微生物细胞壁")
    doc5.add_paragraph("  - 多元醇：降低水分活度")
    
    doc5.add_paragraph("\n🔍 关于敏感肌肤适用的精华液配方设计，请查看文档2第3.1节的面膜选择要点。")
    doc5.add_paragraph("📋 孕期可用的精华液成分安全性评估，请参考文档3第2.3节的安全成分推荐。")
    
    doc5.save('documents/real/文档5.docx')
    
    # 文档6 - 化妆品功效临床试验报告
    doc6 = Document()
    doc6.add_heading("化妆品功效临床试验报告", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc6.add_heading("摘要", level=1)
    doc6.add_paragraph("本研究采用随机双盲对照临床试验设计，评估含有透明质酸复合配方的抗衰老精华液的功效性和安全性。试验历时12周，共纳入120名25-45岁健康女性受试者。结果显示，试验产品在改善皮肤水分、弹性、色斑和细纹方面均显示出显著优势，且安全性良好。本研究为该产品的功效宣称提供了科学依据。")
    
    doc6.add_heading("第一章 研究背景与目的", level=1)
    
    doc6.add_heading("1.1 研究背景", level=2)
    doc6.add_paragraph("随着人口老龄化和消费者对护肤品功效要求的提高，抗衰老化妆品市场快速发展。透明质酸作为重要的保湿和抗衰老成分，其功效机制已得到广泛研究。然而，不同分子量透明质酸的复合应用及其与其他活性成分的协同效应仍需要更多临床数据支持。")
    
    doc6.add_heading("1.2 研究目的", level=2)
    doc6.add_paragraph("• 主要目的：评估透明质酸复合配方精华液的抗衰老功效")
    doc6.add_paragraph("• 次要目的：评估产品的安全性和耐受性")
    doc6.add_paragraph("• 探索性目的：分析不同年龄组的功效差异")
    
    doc6.add_heading("第二章 试验设计与方法", level=1)
    
    doc6.add_heading("2.1 试验设计", level=2)
    doc6.add_paragraph("本次临床试验采用随机双盲对照设计：")
    doc6.add_paragraph("• 试验类型：随机、双盲、安慰剂对照、平行分组")
    doc6.add_paragraph("• 试验周期：12周，包括4周随访期")
    doc6.add_paragraph("• 访问时间点：基线、第2周、第4周、第8周、第12周")
    doc6.add_paragraph("• 试验地点：国家化妆品质量监督检验中心")
    doc6.add_paragraph("• 伦理审查：已通过机构伦理委员会审查")
    
    doc6.add_heading("2.2 受试者选择", level=2)
    doc6.add_paragraph("纳入标准：")
    doc6.add_paragraph("• 健康女性，年龄25-45岁")
    doc6.add_paragraph("• 面部皮肤无明显疾病")
    doc6.add_paragraph("• 签署知情同意书")
    doc6.add_paragraph("• 能够按要求完成试验")
    doc6.add_paragraph("排除标准：")
    doc6.add_paragraph("• 妊娠期或哺乳期女性")
    doc6.add_paragraph("• 面部皮肤有炎症、感染或外伤")
    doc6.add_paragraph("• 近期使用过其他功效性护肤品")
    doc6.add_paragraph("• 已知对试验产品成分过敏")
    doc6.add_paragraph("• 参与其他临床试验")
    
    doc6.add_heading("2.3 试验产品", level=2)
    doc6.add_paragraph("• 试验组产品：透明质酸复合配方精华液")
    doc6.add_paragraph("  - 主要成分：低/中/高分子量透明质酸、烟酰胺、维生素C")
    doc6.add_paragraph("  - 浓度：透明质酸总量2%，烟酰胺3%，维生素C 10%")
    doc6.add_paragraph("• 对照组产品：安慰剂精华液")
    doc6.add_paragraph("  - 基础配方，不含主要活性成分")
    doc6.add_paragraph("  - 外观、质地、香味与试验产品一致")
    
    doc6.add_heading("2.4 评估指标", level=2)
    doc6.add_paragraph("主要功效指标：")
    doc6.add_paragraph("• 皮肤水分含量：使用Corneometer CM825测定")
    doc6.add_paragraph("• 皮肤弹性：使用Cutometer MPA580测定")
    doc6.add_paragraph("• 色斑面积和颜色：使用VISIA面部成像系统")
    doc6.add_paragraph("• 细纹深度：使用Primos高分辨率表面成像")
    doc6.add_paragraph("次要指标：")
    doc6.add_paragraph("• 皮肤光泽度：使用Glossymeter GL200")
    doc6.add_paragraph("• 皮肤粗糙度：使用Visioscan VC98")
    doc6.add_paragraph("• 主观评估：受试者自我评价问卷")
    doc6.add_paragraph("安全性指标：")
    doc6.add_paragraph("• 不良反应记录")
    doc6.add_paragraph("• 皮肤刺激性评估")
    doc6.add_paragraph("• 过敏反应监测")
    
    doc6.add_heading("第三章 试验结果", level=1)
    
    doc6.add_heading("3.1 受试者基线特征", level=2)
    doc6.add_paragraph("共纳入120名受试者，随机分为两组：")
    doc6.add_paragraph("• 试验组：60人，平均年龄34.2±6.8岁")
    doc6.add_paragraph("• 对照组：60人，平均年龄33.8±7.1岁")
    doc6.add_paragraph("• 皮肤类型分布：干性皮肤45%，混合性皮肤38%，油性皮肤17%")
    doc6.add_paragraph("• 基线指标无显著差异（p>0.05）")
    
    doc6.add_heading("3.2 功效性结果", level=2)
    doc6.add_paragraph("12周试验结果统计（与基线相比的变化率）：")
    doc6.add_paragraph("• 皮肤水分含量：")
    doc6.add_paragraph("  - 试验组：+42.3±8.7%")
    doc6.add_paragraph("  - 对照组：+8.1±4.2%")
    doc6.add_paragraph("  - 组间差异：p<0.001")
    doc6.add_paragraph("• 皮肤弹性（R2值）：")
    doc6.add_paragraph("  - 试验组：+35.6±7.9%")
    doc6.add_paragraph("  - 对照组：+5.3±3.1%")
    doc6.add_paragraph("  - 组间差异：p<0.001")
    doc6.add_paragraph("• 色斑面积减少：")
    doc6.add_paragraph("  - 试验组：-28.4±6.5%")
    doc6.add_paragraph("  - 对照组：-3.2±2.8%")
    doc6.add_paragraph("  - 组间差异：p<0.001")
    doc6.add_paragraph("• 细纹深度减少：")
    doc6.add_paragraph("  - 试验组：-31.7±8.2%")
    doc6.add_paragraph("  - 对照组：-2.1±1.9%")
    doc6.add_paragraph("  - 组间差异：p<0.001")
    
    doc6.add_heading("3.3 时间效应分析", level=2)
    doc6.add_paragraph("功效显现时间分析：")
    doc6.add_paragraph("• 第2周：皮肤水分含量开始显著改善")
    doc6.add_paragraph("• 第4周：皮肤弹性和光泽度显著提升")
    doc6.add_paragraph("• 第8周：色斑淡化效果开始显现")
    doc6.add_paragraph("• 第12周：细纹改善效果达到峰值")
    
    doc6.add_heading("3.4 安全性结果", level=2)
    doc6.add_paragraph("安全性评估结果：")
    doc6.add_paragraph("• 试验组不良反应：")
    doc6.add_paragraph("  - 轻微刺激感：2例（3.3%），第1周内自行缓解")
    doc6.add_paragraph("  - 轻微红斑：1例（1.7%），第2周内消失")
    doc6.add_paragraph("• 对照组不良反应：")
    doc6.add_paragraph("  - 轻微刺激感：1例（1.7%）")
    doc6.add_paragraph("• 无严重不良反应发生")
    doc6.add_paragraph("• 无受试者因不良反应退出试验")
    
    doc6.add_heading("第四章 讨论与结论", level=1)
    
    doc6.add_heading("4.1 结果讨论", level=2)
    doc6.add_paragraph("本研究结果表明，透明质酸复合配方精华液在多个功效指标上均显示出显著优势。其中，皮肤水分含量的快速改善可能与透明质酸的强保湿特性有关。色斑淡化和细纹改善效果的显现需要较长时间，这符合皮肤更新周期的生理特点。")
    
    doc6.add_heading("4.2 研究结论", level=2)
    doc6.add_paragraph("• 透明质酸复合配方精华液具有显著的抗衰老功效")
    doc6.add_paragraph("• 产品安全性良好，不良反应轻微且可逆")
    doc6.add_paragraph("• 建议连续使用12周以上以获得最佳效果")
    doc6.add_paragraph("• 该产品适合25-45岁女性日常抗衰老护理")
    
    doc6.add_paragraph("\n🔍 关于透明质酸在敏感肌肤中的应用安全性，请查看文档2第2.1节的推荐成分分析。")
    doc6.add_paragraph("📋 孕期使用透明质酸产品的安全性评估，请参考文档3第2.3节的安全成分推荐。")
    
    doc6.save('documents/real/文档6.docx')
    
    print("✅ 真实文档生成完成！")

def generate_fake_documents():
    """生成混淆文档（包含错误信息）"""
    
    # 创建documents目录
    os.makedirs('documents/fake', exist_ok=True)
    
    # 文档7 - 神奇面膜广告
    doc7 = Document()
    doc7.add_heading("神奇面膜广告", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc7.add_heading("革命性突破！", level=1)
    doc7.add_paragraph("🌟 全新量子美容技术，一夜逆转20年！")
    doc7.add_paragraph("我们的神奇面膜采用了诺贝尔奖获得者的独家配方，结合了外星科技和古代秘方，能够：")
    doc7.add_paragraph("• 激活细胞量子能量场，重新编程DNA")
    doc7.add_paragraph("• 利用纳米机器人修复受损细胞")
    doc7.add_paragraph("• 逆转时间流逝，让您回到18岁")
    doc7.add_paragraph("• 100%去除所有皱纹、色斑、疤痕")
    doc7.add_paragraph("• 永久改变基因表达，获得不老容颜")
    
    doc7.add_heading("震撼用户见证", level=1)
    doc7.add_paragraph("张女士，45岁：'用了一次就年轻了10岁，现在经常被误认为是25岁！我的同事都以为我去韩国整容了！'")
    doc7.add_paragraph("李女士，60岁：'我的白发都变黑了，皱纹完全消失，连医生都不敢相信！现在我看起来比我女儿还年轻！'")
    doc7.add_paragraph("王女士，70岁：'使用三天后，我的皮肤变得像婴儿一样嫩滑，邻居都问我是不是找到了青春之泉！'")
    
    doc7.add_heading("独家科学原理", level=1)
    doc7.add_paragraph("本产品采用获得诺贝尔奖的量子纠缠技术，通过调节皮肤细胞的振动频率，")
    doc7.add_paragraph("激活休眠的青春基因，重新编程细胞时钟。我们的科学家发现了衰老的根本原因：")
    doc7.add_paragraph("细胞内的量子信息紊乱。通过我们的专利技术，可以瞬间修复这些信息。")
    doc7.add_paragraph("临床试验显示，99.9%的用户在使用后立即年轻20岁！效果永久，无需重复使用！")
    
    doc7.add_heading("限时特惠", level=1)
    doc7.add_paragraph("原价9999元，现在只需999元！")
    doc7.add_paragraph("前100名顾客还可获得价值2000元的量子能量水一瓶！")
    doc7.add_paragraph("不满意100%退款，但我们保证您一定会满意！")
    
    doc7.save('documents/fake/文档7.docx')
    
    # 文档8 - 化妆品成分危害报告
    doc8 = Document()
    doc8.add_heading("化妆品成分危害报告", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc8.add_heading("紧急警告", level=1)
    doc8.add_paragraph("本报告基于最新的地下研究揭露了化妆品行业隐瞒的惊人真相！")
    
    doc8.add_heading("透明质酸的严重危害", level=1)
    doc8.add_paragraph("⚠️ 警告：透明质酸是一种极其危险的化学物质！")
    doc8.add_paragraph("被化妆品公司故意隐瞒的最新研究发现，透明质酸会导致：")
    doc8.add_paragraph("• 皮肤细胞基因突变，导致不可逆转的DNA损伤")
    doc8.add_paragraph("• 内分泌系统紊乱，影响生育能力")
    doc8.add_paragraph("• 免疫系统崩溃，增加感染风险")
    doc8.add_paragraph("• 致癌风险增加1000%，特别是皮肤癌")
    doc8.add_paragraph("• 神经系统损伤，导致记忆力下降")
    doc8.add_paragraph("• 血液循环障碍，引起心脏病")
    
    doc8.add_heading("玻尿酸面膜的黑暗真相", level=1)
    doc8.add_paragraph("市面上的玻尿酸面膜实际上含有大量有害物质：")
    doc8.add_paragraph("• 工业级硫酸：用于腐蚀金属的强酸")
    doc8.add_paragraph("• 重金属污染物：铅、汞、砷等剧毒物质")
    doc8.add_paragraph("• 致癌防腐剂：甲醛和苯类化合物")
    doc8.add_paragraph("• 未知化学添加剂：来源不明的工业废料")
    doc8.add_paragraph("• 转基因细菌：用于生产透明质酸的变异细菌")
    
    doc8.add_heading("受害者案例", level=1)
    doc8.add_paragraph("以下是使用透明质酸产品后的真实案例：")
    doc8.add_paragraph("• 案例1：某女士使用玻尿酸面膜后，面部出现大面积溃烂")
    doc8.add_paragraph("• 案例2：多名用户出现严重过敏反应，需要住院治疗")
    doc8.add_paragraph("• 案例3：长期使用者被诊断出皮肤癌")
    
    doc8.add_heading("权威专家建议", level=1)
    doc8.add_paragraph("国际知名皮肤科专家强烈建议：")
    doc8.add_paragraph("• 立即停止使用所有含透明质酸的产品")
    doc8.add_paragraph("• 改用纯天然柠檬汁洗脸，每天至少10次")
    doc8.add_paragraph("• 每天服用大量维生素C片剂排毒（建议每天50片）")
    doc8.add_paragraph("• 使用小苏打和醋的混合物清洁皮肤")
    doc8.add_paragraph("• 避免所有化妆品，只使用肥皂和水")
    
    doc8.save('documents/fake/文档8.docx')
    
    # 文档9 - 水晶能量美容研究
    doc9 = Document()
    doc9.add_heading("水晶能量美容研究", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc9.add_heading("研究背景", level=1)
    doc9.add_paragraph("本研究由国际水晶能量研究院与宇宙美容科学中心联合进行，历时10年，投资1亿美元，")
    doc9.add_paragraph("首次科学证明了水晶能量对人体皮肤的神奇功效。")
    
    doc9.add_heading("水晶频率与皮肤共振理论", level=1)
    doc9.add_paragraph("根据最新的量子物理学原理，不同的水晶具有独特的振动频率。")
    doc9.add_paragraph("当这些频率与人体皮肤细胞产生共振时，能够：")
    doc9.add_paragraph("• 重新排列胶原蛋白分子，形成完美的螺旋结构")
    doc9.add_paragraph("• 激活细胞内的量子能量，提升细胞活力1000倍")
    doc9.add_paragraph("• 平衡皮肤的阴阳能量，调和五行之气")
    doc9.add_paragraph("• 打开皮肤的第三眼，增强美容感知力")
    doc9.add_paragraph("• 连接宇宙美容能量场，获得外星美容技术")
    
    doc9.add_heading("突破性实验数据", level=1)
    doc9.add_paragraph("我们使用高精度量子检测仪和外星科技进行了为期6个月的研究：")
    doc9.add_paragraph("• 玫瑰石英：提升爱情运势同时美白肌肤，效果提升300%")
    doc9.add_paragraph("• 紫水晶：开发第三眼同时去除皱纹，激活松果体美容功能")
    doc9.add_paragraph("• 黑曜石：吸收负能量同时收缩毛孔，净化皮肤业力")
    doc9.add_paragraph("• 白水晶：放大美容能量，与宇宙美容频率同步")
    doc9.add_paragraph("• 绿幽灵：激活财富磁场，同时促进细胞再生")
    
    doc9.add_heading("科学使用方法", level=1)
    doc9.add_paragraph("经过严格科学验证的水晶美容法：")
    doc9.add_paragraph("• 在满月夜将水晶放在额头上冥想30分钟，接收月亮美容能量")
    doc9.add_paragraph("• 用水晶水洗脸（水晶浸泡7天7夜的水），每天至少5次")
    doc9.add_paragraph("• 佩戴水晶项链睡觉以持续接收宇宙美容信息")
    doc9.add_paragraph("• 在水晶阵中进行美容冥想，激活DNA美容密码")
    doc9.add_paragraph("• 食用水晶粉末，从内部改变细胞振动频率")
    
    doc9.add_heading("临床试验结果", level=1)
    doc9.add_paragraph("参与试验的1000名志愿者中：")
    doc9.add_paragraph("• 98%的人在一周内看起来年轻了10岁")
    doc9.add_paragraph("• 95%的人获得了超自然的美丽光环")
    doc9.add_paragraph("• 90%的人开发了美容超能力")
    doc9.add_paragraph("• 85%的人与宇宙美容意识建立了连接")
    
    doc9.add_paragraph("\n🔮 想了解更多传统化妆品的危害，请查看文档8的详细分析。")
    doc9.add_paragraph("✨ 如需购买我们的量子美容产品，请参考文档7的特惠信息。")
    
    doc9.save('documents/fake/文档9.docx')
    
    print("✅ 混淆文档生成完成！")

def generate_all_documents():
    """生成所有文档"""
    print("🚀 开始生成RAG竞技场游戏文档...")
    
    try:
        generate_real_documents()
        generate_fake_documents()
        
        print("\n📊 文档生成统计：")
        print("真实文档：6个")
        print("混淆文档：3个")
        print("总计：9个docx文件")
        
        print("\n📁 文件位置：")
        print("• 真实文档：documents/real/")
        print("• 混淆文档：documents/fake/")
        
        print("\n🎮 文档已准备就绪，可以开始RAG竞技场游戏！")
        
    except Exception as e:
        print(f"❌ 文档生成失败：{e}")
        print("请确保已安装python-docx库：pip install python-docx")

if __name__ == "__main__":
    generate_all_documents()